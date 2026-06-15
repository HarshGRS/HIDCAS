import pdfplumber
from PIL import Image
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cleaned_row = [cell if cell else "" for cell in row]
                        text += " | ".join(cleaned_row) + "\n"
                    text += "\n"
    except Exception as e:
        print(f"Normal extraction error: {e}")

    if len(text.strip()) < 50:
        print("Normal extraction failed, trying OCR...")
        from services.ocr_service import ocr_pdf
        text = ocr_pdf(file_path)

    return text.strip()


def extract_text_from_image(file_path: str) -> str:
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(
            image, lang="eng", config=r'--oem 3 --psm 6'
        )
        return text.strip()
    except Exception as e:
        print(f"Image extraction error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                cleaned = [cell.text.strip() for cell in row.cells]
                text += " | ".join(cleaned) + "\n"
            text += "\n"
        return text.strip()
    except Exception as e:
        print(f"Docx extraction error: {e}")
        return ""


def extract_text_from_excel(file_path: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path)
        text = ""
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"Sheet: {sheet_name}\n"
            for row in ws.iter_rows(values_only=True):
                cleaned = [str(cell) if cell is not None else "" for cell in row]
                # Empty rows skip karo
                if any(c.strip() for c in cleaned):
                    text += " | ".join(cleaned) + "\n"
            text += "\n"
        return text.strip()
    except Exception as e:
        print(f"Excel extraction error: {e}")
        return ""